#!/usr/bin/env python3
"""Update Section 14 White Balance — rename presets, fix preview column, update text."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r'C:\Users\tomcw\App_tcw3\iCamera\iCamera_User_Manual.docx'
doc = Document(DOC_PATH)

def find_para(fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1

def replace_text(para, new_text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

# ─────────────────────────────────────────────────────────────────────────────
# 1. Rebuild the WB presets table with correct names and corrected descriptions
# ─────────────────────────────────────────────────────────────────────────────
new_rows = [
    ['Kelvin', 'Preset Name', 'Light source it corrects for', 'Effect on preview & saved photo'],
    ['2800 K', 'Bulb',
     'Incandescent light bulb — warm orange/yellow indoor light',
     'Cuts red, boosts blue — removes yellow cast, whites look white'],
    ['3500 K', 'Indoor',
     'Warm LED or halogen bulb',
     'Moderate cool shift — reduces warm tint'],
    ['4500 K', 'Fluorescent',
     'Office fluorescent tubes',
     'Slight cool shift — removes greenish/warm tint'],
    ['5500 K', 'Daylight  ★',
     'Midday sun — the neutral default',
     'No correction — true neutral, no colour cast'],
    ['6500 K', 'Cloudy',
     'Overcast or cloudy sky — slightly cool bluish light',
     'Slight warm shift — removes blue cast'],
    ['8000 K', 'Shade',
     'Open shade or heavily overcast — very cool blue light',
     'Boosts red, cuts blue — removes blue cast, adds warmth'],
]

# Find and replace the existing table in section 14
found_table = False
for block in doc.element.body:
    tag = block.tag.split('}')[-1]
    if tag == 'tbl' and not found_table:
        t = Table(block, doc)
        # Check if it's the WB table (first cell = "Kelvin")
        if t.rows[0].cells[0].text.strip() == 'Kelvin':
            found_table = True
            # Rebuild all rows
            for r_idx, row_data in enumerate(new_rows):
                if r_idx < len(t.rows):
                    row = t.rows[r_idx]
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = cell_text
                            if r_idx == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            print('Table: Updated WB presets table')

# ─────────────────────────────────────────────────────────────────────────────
# 2. Update the footnote below the table
# ─────────────────────────────────────────────────────────────────────────────
i = find_para('★ 5500 K Daylight is the default')
if i >= 0:
    replace_text(doc.paragraphs[i],
        '★ 5500 K Daylight is the default and is correct for most outdoor shooting in sunlight.')
    print('Updated table footnote')

# ─────────────────────────────────────────────────────────────────────────────
# 3. Update "Does White Balance Affect the Saved Photo?" section
# ─────────────────────────────────────────────────────────────────────────────
i = find_para('Does White Balance Affect the Saved Photo')
if i >= 0:
    # Update the heading
    replace_text(doc.paragraphs[i], 'Does White Balance Affect the Preview and Saved Photo?')
    print('Updated section heading')

i = find_para('Yes — the WB correction is applied to every saved photo')
if i >= 0:
    replace_text(doc.paragraphs[i],
        'Yes — to both. The moment you tap a new WB preset, the live viewfinder '
        'immediately shifts colour to show you the correction. '
        'The same correction is then baked into every saved photo in all shooting modes '
        '(AUTO, PRO, and APT). '
        'What you see in the preview is exactly what you get in the saved photo.')
    print('Updated "affects preview and photo" paragraph')

# ─────────────────────────────────────────────────────────────────────────────
# 4. Update the Practical Tips bullets to use new preset names
# ─────────────────────────────────────────────────────────────────────────────
replacements = {
    'Indoors under warm bulbs → set WB to 2800 K or 3500 K':
        'Indoors under a light bulb → tap WB and set to Bulb (2800 K) or Indoor (3500 K). '
        'The yellow cast disappears instantly in the preview — what you see is what you get.',
    'Outdoors on a cloudy day → 6500 K or 8000 K':
        'Outdoors on a cloudy or overcast day → set WB to Cloudy (6500 K). '
        'This removes the cool blue cast that overcast light creates.',
    'Outdoors in bright sun → leave at 5500 K':
        'Outdoors in bright sun → leave at Daylight (5500 K). This is the neutral default.',
    'Creative use: deliberately set the wrong WB':
        'Creative use: deliberately set the "wrong" WB for effect. '
        'Bulb (2800 K) outdoors gives a dramatic cool/cinematic blue look. '
        'Shade (8000 K) indoors adds a golden warm-hour feel.',
    'The WB preview and the saved photo always match':
        'Preview = saved photo. The colour correction applied in the viewfinder is '
        'identical to what is baked into the saved image — no surprises.',
}

for old_fragment, new_text in replacements.items():
    i = find_para(old_fragment)
    if i >= 0:
        replace_text(doc.paragraphs[i], new_text)
        print(f'Updated bullet: {old_fragment[:50]}...')

# ─────────────────────────────────────────────────────────────────────────────
# 5. Update the TOC section 3 WB bullet reference
# ─────────────────────────────────────────────────────────────────────────────
i = find_para('WB — White Balance in Kelvin')
if i >= 0:
    replace_text(doc.paragraphs[i],
        'WB — White Balance. Tap to cycle presets: Bulb / Indoor / Fluorescent / '
        'Daylight / Cloudy / Shade. '
        'The viewfinder shifts colour immediately. See Section 14.')
    print('Updated HUD WB bullet in Section 3')

doc.save(DOC_PATH)
print('\nManual saved.')
