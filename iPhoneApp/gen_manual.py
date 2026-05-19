"""
iCamera User Manual generator — combines existing manual content with new
screenshots, updated UI descriptions, and additional reference sections.
Output: iCamera_User_Manual_v2.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

BASE = r"C:\Users\tomcw\App_tcw3\iCamera"
SS   = os.path.join(BASE, "docs", "screenshots")

# ── colour palette ──────────────────────────────────────────────────────────
C_LEICA_RED  = RGBColor(0xC8, 0x00, 0x00)   # Leica accent red
C_DARK       = RGBColor(0x1A, 0x1A, 0x1A)   # near-black text
C_MID        = RGBColor(0x44, 0x44, 0x44)   # body text
C_SUBTLE     = RGBColor(0x66, 0x66, 0x66)   # captions / labels
C_RULE       = RGBColor(0xCC, 0xCC, 0xCC)   # horizontal rules
C_TBL_HDR    = RGBColor(0x22, 0x22, 0x22)   # table header bg
C_TBL_ALT    = RGBColor(0xF5, 0xF5, 0xF5)   # table alt row bg
C_ACCENT     = RGBColor(0xC8, 0x00, 0x00)   # heading accent
FONT_BODY    = "Georgia"
FONT_HEADING = "Georgia"
FONT_MONO    = "Courier New"

doc = Document()

# ── page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── style helpers ───────────────────────────────────────────────────────────
def set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'),     str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc, color="CCCCCC", thickness=6):
    para = doc.add_paragraph()
    set_para_spacing(para, before=60, after=60)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    str(thickness))
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), color)
    pBdr.append(btm)
    pPr.append(pBdr)
    return para

def heading1(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=240, after=80)
    run  = para.add_run(text.upper())
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(14)
    run.font.bold      = True
    run.font.color.rgb = C_LEICA_RED
    run.font.color.theme_color = None
    # bottom border accent line
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '4')
    btm.set(qn('w:space'), '4')
    btm.set(qn('w:color'), 'C80000')
    pBdr.append(btm)
    pPr.append(pBdr)
    return para

def heading2(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=160, after=40)
    run  = para.add_run(text)
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.color.rgb = C_DARK
    return para

def heading3(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=100, after=30)
    run  = para.add_run(text)
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.italic    = True
    run.font.color.rgb = C_MID
    return para

def body(doc, text, indent=False):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=60, line=276)
    if indent:
        para.paragraph_format.left_indent = Cm(0.6)
    run  = para.add_run(text)
    run.font.name  = FONT_BODY
    run.font.size  = Pt(10)
    run.font.color.rgb = C_MID
    return para

def bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    set_para_spacing(para, before=0, after=40, line=260)
    para.paragraph_format.left_indent   = Cm(0.8)
    para.paragraph_format.first_line_indent = Cm(-0.4)
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = C_DARK
        r = para.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(10)
        r.font.color.rgb = C_MID
    else:
        r = para.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(10)
        r.font.color.rgb = C_MID
    return para

def caption(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=30, after=80)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run(text)
    run.font.name   = FONT_BODY
    run.font.size   = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = C_SUBTLE
    return para

def add_image(doc, path, width_cm=14, align=WD_ALIGN_PARAGRAPH.CENTER):
    para = doc.add_paragraph()
    set_para_spacing(para, before=60, after=30)
    para.alignment = align
    run  = para.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return para

def tip_box(doc, text):
    """Shaded tip paragraph."""
    para = doc.add_paragraph()
    set_para_spacing(para, before=60, after=60)
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FFF3F3')
    pPr.append(shd)
    r = para.add_run("▶  TIP:  ")
    r.font.name      = FONT_BODY
    r.font.size      = Pt(9.5)
    r.font.bold      = True
    r.font.color.rgb = C_LEICA_RED
    r2 = para.add_run(text)
    r2.font.name     = FONT_BODY
    r2.font.size     = Pt(9.5)
    r2.font.color.rgb = C_MID
    return para

def two_col_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, '222222')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        set_para_spacing(para, before=40, after=40)
        run  = para.add_run(h)
        run.font.name      = FONT_BODY
        run.font.size      = Pt(9)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data rows
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            set_para_spacing(para, before=30, after=30)
            # first column bold
            run  = para.add_run(val)
            run.font.name      = FONT_BODY
            run.font.size      = Pt(9)
            run.font.bold      = (ci == 0)
            run.font.color.rgb = C_DARK if ci == 0 else C_MID
    # column widths
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return tbl

def add_spacer(doc, pts=6):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=pts)
    para.add_run("")
    return para

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=240, after=80)
run = para.add_run()
run.add_picture(os.path.join(SS, "icon.png"), width=Cm(4.5))

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=60, after=20)
run  = para.add_run("i Camera")
run.font.name      = FONT_HEADING
run.font.size      = Pt(40)
run.font.bold      = True
run.font.color.rgb = C_DARK

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=0, after=20)
run  = para.add_run("Professional Camera App")
run.font.name      = FONT_HEADING
run.font.size      = Pt(16)
run.font.color.rgb = C_SUBTLE

add_horizontal_rule(doc, color="C80000", thickness=8)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=40, after=200)
run  = para.add_run("Inspired by Leica LUX  ·  Full Manual Control  ·  Film-Grade Colour Science")
run.font.name      = FONT_BODY
run.font.size      = Pt(11)
run.font.italic    = True
run.font.color.rgb = C_SUBTLE

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=0, after=12)
run  = para.add_run("User Manual  ·  Version 1.0  ·  May 2026  ·  Android")
run.font.name      = FONT_BODY
run.font.size      = Pt(10)
run.font.color.rgb = C_SUBTLE

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=0, after=0)
run  = para.add_run("tcw3  ·  tomcwong3@gmail.com")
run.font.name      = FONT_BODY
run.font.size      = Pt(9)
run.font.color.rgb = C_RULE

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (manual — Word auto-TOC requires field codes)
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "Table of Contents")

toc_items = [
    ("1.",  "Introduction"),
    ("2.",  "Getting Started"),
    ("3.",  "Camera Screen Overview"),
    ("4.",  "Shooting Modes — AUTO, PRO, APT"),
    ("5.",  "PRO Mode: ISO & Shutter Speed Explained"),
    ("6.",  "APT Mode: Aperture & Bokeh"),
    ("7.",  "Film Looks (Look Filters)"),
    ("8.",  "Lens Profiles — 28 mm / 35 mm / 50 mm"),
    ("9.",  "White Balance"),
    ("10.", "Flash"),
    ("11.", "RAW Capture"),
    ("12.", "Photo Quality — STD vs HQ"),
    ("13.", "Saving & Finding Your Photos"),
    ("14.", "Tips & Troubleshooting"),
]
for num, title in toc_items:
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=40)
    r1 = para.add_run(f"{num:<5}")
    r1.font.name      = FONT_BODY
    r1.font.size      = Pt(10)
    r1.font.bold      = True
    r1.font.color.rgb = C_LEICA_RED
    r2 = para.add_run(title)
    r2.font.name      = FONT_BODY
    r2.font.size      = Pt(10)
    r2.font.color.rgb = C_MID

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "1.  Introduction")
body(doc,
     "iCamera is a professional camera app for Android, designed with the clean aesthetic of Leica "
     "rangefinder cameras. It gives you full manual control over your phone camera — ISO, shutter "
     "speed, aperture simulation, lens profiles, and cinematic colour looks — all in one place.")
body(doc,
     "Built on a native C++ image pipeline with real-time 3D LUT processing, iCamera applies film "
     "emulation, optical lens simulation, and AI-assisted depth effects entirely on-device. No cloud "
     "upload, no subscription required.")
body(doc,
     "This manual explains every feature in plain English. No photography degree required.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "2.  Getting Started")

heading2(doc, "Permissions")
body(doc, "When you launch iCamera for the first time the app will request two permissions:")
bullet(doc, "Camera — required to take photos.")
bullet(doc, "Storage / Media — required to save photos to your gallery.")
body(doc,
     "Tap Allow on both prompts. If you accidentally denied them, go to "
     "Settings → Apps → iCamera → Permissions and enable them manually.")

heading2(doc, "First Launch")
body(doc,
     "The app opens directly to the camera viewfinder. Point your phone at a subject and tap the "
     "large round shutter button at the bottom to take a photo. Your photo is automatically saved "
     "to Pictures/iCamera/ on your device.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3. CAMERA SCREEN OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "3.  Camera Screen Overview")

# Portrait screenshot
add_image(doc, os.path.join(SS, "portrait_pro.jpg"), width_cm=8)
caption(doc, "Figure 1 — Portrait mode, PRO active. Film look selector, WB row, and manual exposure wheels visible.")

heading2(doc, "Top Bar (HUD — Heads-Up Display)")
body(doc, "The row of values across the top of the screen shows your current settings at a glance:")
bullet(doc, "Shutter Speed (e.g. 1/2000) — how long the sensor collects light.", bold_prefix="SS  ")
bullet(doc, "Sensor sensitivity (e.g. ISO 400).", bold_prefix="ISO  ")
bullet(doc, "Active virtual lens name (centre display, e.g. NOCTILUX-M 50).", bold_prefix="Lens  ")
bullet(doc, "White Balance colour temperature (e.g. 5500K Daylight).", bold_prefix="WB  ")
bullet(doc, "Simulated aperture (e.g. f/5.6) — controls bokeh depth.", bold_prefix="APT  ")

body(doc,
     "The film look name (e.g. CLASSIC) is displayed as a badge in the top-right corner so you "
     "always know which emulation is active while composing.")

heading2(doc, "Bottom Controls — Portrait Mode")
body(doc, "The bottom panel is organised into rows from top to bottom:")

heading3(doc, "Row 1 — Film Look  +  Lens Selector")
body(doc,
     "Film:  CLASSIC   CONTEMPORARY   B&W   VIVID   ARTISTIC  "
     "                         28mm   35mm   50mm")
bullet(doc,
       "Scroll horizontally to choose the film emulation. The look is applied at capture time via "
       "a full 3D Look-Up Table.",
       bold_prefix="Film: row  ")
bullet(doc,
       "Tap 28mm, 35mm, or 50mm to switch the active lens profile. Each lens snaps the zoom to "
       "its default level and changes vignette, aberration, and distortion.",
       bold_prefix="Lens buttons  ")

heading3(doc, "Row 2 — White Balance")
body(doc, "WB:   BULB   INDOOR   FLUORESCENT   DAYLIGHT   CLOUDY   SHADE")
body(doc, "Tap any preset to shift the colour temperature of the captured image.")

heading3(doc, "Exposure Wheels (PRO Mode Only)")
body(doc,
     "In PRO mode, two horizontal scroll wheels appear — SHUTTER on the left and ISO on the right. "
     "The selected value is always centred above the red tick mark. Scroll left or right and lift "
     "your finger to apply.")

heading3(doc, "Mode Bar")
body(doc, "RAW  |  AUTO  |  PRO  |  APT  |  FLASH  |  FLIP  |  STD/HQ")
bullet(doc, "Toggle DNG capture. Film pipeline is bypassed; raw sensor data is saved.", bold_prefix="RAW  ")
bullet(doc, "Fully automatic exposure.", bold_prefix="AUTO  ")
bullet(doc, "Full manual ISO + shutter speed.", bold_prefix="PRO  ")
bullet(doc, "Aperture-priority — you set f-stop, camera sets shutter.", bold_prefix="APT  ")
bullet(doc, "Enable / disable flash.", bold_prefix="FLASH  ")
bullet(doc, "Switch between rear and front camera.", bold_prefix="FLIP  ")
bullet(doc, "Toggle image quality between Standard (~8MP) and High Quality (13MP).", bold_prefix="STD/HQ  ")

heading3(doc, "Shutter Button & Thumbnail")
body(doc,
     "The large circle at the bottom centre is the shutter button. Tap once to capture. "
     "The shutter releases immediately — film processing runs in the background so you are "
     "ready for the next shot without waiting. The small square in the bottom-left corner "
     "shows a thumbnail of the last captured photo. Tap it to open the system gallery.")

add_spacer(doc)

# Landscape screenshot
add_image(doc, os.path.join(SS, "landscape_auto.jpg"), width_cm=14)
caption(doc, "Figure 2 — Landscape mode, AUTO active. Full-screen preview with HUD and mode controls.")

add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 4. SHOOTING MODES
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "4.  Shooting Modes — AUTO, PRO, APT")

heading2(doc, "AUTO Mode")
body(doc,
     "The phone's automatic camera brain handles everything — brightness, focus, and colour balance. "
     "This is the fastest mode for everyday shooting. Just point and shoot. Film looks, white balance, "
     "and lens simulation still apply.")

heading2(doc, "PRO Mode")
body(doc,
     "You take full manual control of ISO and Shutter Speed. The automatic exposure brain is switched "
     "off. Two scroll wheels appear above the mode bar — SHUTTER on the left and ISO on the right. "
     "Scroll left or right to change values. Lift your finger to apply.")
tip_box(doc,
        "Set ISO 100 and shutter 1/500 in bright daylight. In a dim interior try ISO 800 and "
        "shutter 1/60. See Section 5 for a full plain-English explanation.")

heading2(doc, "APT Mode (Aperture)")
body(doc,
     "Simulates a real camera lens with a variable aperture ring. A dial lets you set the f-stop "
     "(e.g. f/1.4, f/2.8, f/8). A lower f-number gives a blurred background. A higher f-number "
     "keeps more of the scene sharp.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 5. PRO MODE — ISO & SHUTTER EXPLAINED
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "5.  PRO Mode: ISO & Shutter Speed Explained in Plain English")
body(doc,
     "This section explains what ISO and Shutter Speed actually do — without photography jargon. "
     "Understanding these two controls lets you take better photos in any lighting condition.")

heading3(doc, "In one sentence")
body(doc,
     "Shutter controls how LONG the sensor collects light (time → brightness and motion blur). "
     "ISO controls how LOUDLY the sensor amplifies the light it already collected "
     "(gain → brightness and grain). Both affect brightness — but through different mechanisms.",
     indent=True)

add_horizontal_rule(doc)

heading2(doc, "ISO — How Much the Sensor Amplifies Light")
body(doc,
     "On a traditional film camera, ISO measured the film’s sensitivity. On a phone the sensor "
     "is a fixed silicon chip that cannot become more or less sensitive. Instead it amplifies the "
     "electrical signal that light creates:")
bullet(doc,
       "The chip amplifies the light signal very little. The photo is clean with no grain, but "
       "needs plenty of light or a slower shutter.",
       bold_prefix="LOW ISO (e.g. 100)  ")
bullet(doc,
       "The chip amplifies the same signal much more. The photo is brighter in dark conditions, "
       "but electronic noise creates visible grain. ISO 800 is moderate — acceptable grain. "
       "ISO 6400 is a heavy boost — noticeably grainy.",
       bold_prefix="HIGH ISO (e.g. 800–6400)  ")

body(doc,
     "Think of it like a microphone in a quiet room. ISO is the volume amplifier on the recording. "
     "Low ISO (100) = amplifier at minimum — clean and clear, but needs plenty of light. "
     "High ISO (3200) = amplifier turned up — image gets brighter in dark scenes, but "
     "background electronic hiss also gets louder. That hiss is the grain you see in the photo.")
tip_box(doc, "Always use the lowest ISO that still gives you a bright enough image.")

heading3(doc, "A concrete example")
body(doc, "Two photos of the same object, both with shutter 1/100:")
bullet(doc,
       "Sensor collects light for 1/100 s, amplifies very little. Result: clean, low-noise — "
       "but darker in dim light.",
       bold_prefix="Photo A — ISO 100, 1/100:  ")
bullet(doc,
       "Same collection time, but signal amplified 64× before saving. Result: much brighter "
       "— but with visible grain.",
       bold_prefix="Photo B — ISO 6400, 1/100:  ")
body(doc,
     "Key point: ISO does not change how long the sensor collects light — only shutter does that. "
     "ISO is amplification applied to whatever light was already collected.")

add_horizontal_rule(doc)

heading2(doc, "Shutter Speed — How Long the Sensor Collects Light")
body(doc,
     "A phone uses an electronic shutter — the sensor chip simply starts and stops collecting "
     "light electronically:")
bullet(doc,
       "Sensor collects light for only 1 ms. Fast-moving subjects are frozen sharply. "
       "But less light enters, so the photo is darker.",
       bold_prefix="FAST shutter (e.g. 1/1000)  ")
bullet(doc,
       "Sensor collects light for 66 ms. More light enters — brighter photo. But if the "
       "phone or subject moves during those 66 ms, the photo will be blurry.",
       bold_prefix="SLOW shutter (e.g. 1/15)  ")
tip_box(doc,
        "Use fast shutter (1/500 or higher) for sports and action. "
        "Use slow shutter (1/30 or lower) in dim conditions, but rest the phone on a flat surface.")

add_horizontal_rule(doc)

heading2(doc, "How iCamera Sends These Values to the Hardware")
body(doc,
     "Google provides a low-level interface called Camera2 that lets an app talk directly to the "
     "camera sensor chip and say: “Turn off your automatic exposure brain and use exactly this "
     "amplification (ISO) and exactly this collection window (shutter speed) for every frame.”")
body(doc,
     "iCamera sends those two numbers directly via the Android Camera2 API. On devices where full "
     "manual control is unavailable, iCamera automatically falls back to an EV-offset approach — "
     "biasing the auto-exposure system by a calculated number of stops. The result looks the same "
     "to you, but the underlying mechanism is slightly less precise than true hardware control.")

add_horizontal_rule(doc)

heading2(doc, "Practical Guide to ISO & Shutter Speed")

two_col_table(doc,
    ["Situation", "Recommended Settings"],
    [
        ["Bright outdoor daylight",       "ISO 100  ·  Shutter 1/500 or faster — clean, frozen motion"],
        ["Indoors or overcast day",        "ISO 400–800  ·  Shutter 1/60 — balanced brightness"],
        ["Night or very low light",        "ISO 1600–3200  ·  Shutter 1/30 or slower — rest phone on surface"],
        ["Action / sports",               "ISO 400–800  ·  Shutter 1/1000+ — freeze fast movement"],
        ["Intentional motion blur",        "ISO 100  ·  Shutter 1/15 or slower — tripod recommended"],
    ],
    col_widths=[7, 9.5]
)
add_spacer(doc)

heading3(doc, "The Relationship Between ISO and Shutter Speed")
body(doc,
     "ISO and shutter speed always work together. If you make the shutter faster (less light), "
     "compensate by raising ISO (more amplification). If you lower ISO for a cleaner image, you "
     "may need a slower shutter to let in enough light. Finding the right balance between "
     "brightness, grain, and motion is the art of manual photography.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 6. APT MODE
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "6.  APT Mode: Aperture & Bokeh")
body(doc,
     "APT mode simulates a real camera lens with a variable aperture — the opening that controls "
     "how blurred the background is.")
bullet(doc,
       "Very blurred background. Subject stands out clearly. Best for portraits and close-ups.",
       bold_prefix="Low f-number (e.g. f/1.4)  ")
bullet(doc,
       "Most of the scene stays sharp. Good for landscapes and group shots.",
       bold_prefix="High f-number (e.g. f/8 or f/16)  ")
body(doc,
     "Toggle the BOKEH switch to enable background blur. iCamera uses on-device AI (Google ML Kit "
     "Selfie Segmentation) to detect the subject and apply realistic background blur at capture time.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 7. FILM LOOKS
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "7.  Film Looks (Look Filters)")
body(doc,
     "The Film: selector at the top of the bottom controls lets you choose a colour grade for your "
     "photo. Each look applies a professional 3D colour transformation (33×33×33 LUT) "
     "using trilinear interpolation — processed in C++ at capture time. The look is baked into the "
     "saved photo, not just a preview filter.")

two_col_table(doc,
    ["Look", "Character", "Best For"],
    [
        ["CLASSIC",       "Warm amber midtones, lifted shadows, gentle halation",
                          "Portraits, everyday life, golden-hour scenes"],
        ["CONTEMPORARY",  "Neutral-warm, clean, slightly desaturated",
                          "Street photography, architecture, modern subjects"],
        ["B&W",           "High-contrast monochrome, deep blacks",
                          "Dramatic portraits, texture-rich subjects, fine art"],
        ["VIVID",         "Punchy colour, high saturation, rich contrast",
                          "Travel, landscapes, product shots in good light"],
        ["ARTISTIC",      "Cross-processed colour shift, cinematic grade",
                          "Creative / experimental — emulates bleach-bypass film"],
    ],
    col_widths=[3.5, 6, 6]
)
add_spacer(doc, 12)

body(doc,
     "Start with CLASSIC for most situations. Switch to B&W when you want to emphasise shape, "
     "shadow, and texture over colour. Use VIVID for scenes with strong, saturated colours. "
     "ARTISTIC is best used intentionally for a specific mood.")

# Film look comparison
add_spacer(doc)
heading2(doc, "Film Look Comparison — Same Subject, Two Looks")

# Side-by-side table for images
img_tbl = doc.add_table(rows=2, cols=2)
img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in img_tbl.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBdr = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'none')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
            tcBdr.append(el)
        tcPr.append(tcBdr)

for ci, (img, lbl) in enumerate([
        (os.path.join(SS, "sample_classic.jpg"), "CLASSIC — warm amber, lifted shadows"),
        (os.path.join(SS, "sample_bw.jpg"),      "B&W — high-contrast monochrome"),
    ]):
    cell = img_tbl.rows[0].cells[ci]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img, width=Cm(7))
    cell2 = img_tbl.rows[1].cells[ci]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p2, before=20, after=20)
    r = p2.add_run(lbl)
    r.font.name   = FONT_BODY
    r.font.size   = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = C_SUBTLE

add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 8. LENS PROFILES
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "8.  Lens Profiles — 28 mm / 35 mm / 50 mm")
body(doc,
     "The lens selector at the bottom of the screen lets you choose from three Leica-inspired lens "
     "simulations. Each one changes the zoom level of the viewfinder and applies a unique combination "
     "of vignette, chromatic aberration, and barrel distortion to your saved photos.")

two_col_table(doc,
    ["Lens", "Focal Length", "Aperture", "Default Zoom", "Character"],
    [
        ["Summilux 28mm", "Wide-angle",       "f/1.4",  "1.0×",
         "Mild vignette, slight barrel distortion — shows more of the scene"],
        ["Summilux 35mm", "Standard / normal","f/1.4",  "1.3×",
         "Medium vignette, warm-gold tint — closest to human eye perspective"],
        ["Noctilux 50mm", "Short telephoto",  "f/0.95", "1.8×",
         "Strong vignette, warm amber tint — maximum bokeh, portrait favourite"],
    ],
    col_widths=[3.5, 3.2, 2.2, 2.5, 5]
)
add_spacer(doc, 12)

bullet(doc,
       "Wide-angle perspective. Good for street photography, architecture, and group shots. "
       "This is the phone’s native 1.0× — widest, sharpest view with no digital zoom.",
       bold_prefix="Summilux 28mm  ")
bullet(doc,
       "The “standard” lens — closest to what the human eye sees naturally. Versatile "
       "all-rounder for everyday shooting, travel, and street photography.",
       bold_prefix="Summilux 35mm  ")
bullet(doc,
       "Portrait lens with the fastest virtual aperture (f/0.95). Produces the strongest vignette "
       "and background separation. Ideal for close-up portraits and isolating subjects from "
       "distracting backgrounds.",
       bold_prefix="Noctilux 50mm  ")

heading2(doc, "What the Focal Length Numbers Mean")
body(doc,
     "The “mm” number is borrowed from classic 35mm film cameras and describes the field "
     "of view — how wide or how zoomed-in the image looks. A lower number = wider view. A higher "
     "number = narrower, more zoomed view. The number closest to “what your eyes see "
     "naturally” is around 35mm.")

heading2(doc, "Zoom Behaviour")
body(doc,
     "Tapping a lens automatically snaps the zoom to its correct level (28mm → 1.0×, "
     "35mm → 1.3×, 50mm → 1.8×). A zoom indicator appears briefly at the "
     "top of the viewfinder to confirm. You can also pinch the viewfinder with two fingers at any "
     "time to zoom freely between 1× and 10×. The indicator updates as you pinch.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 9. WHITE BALANCE
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "9.  White Balance")
body(doc,
     "White Balance (WB) tells the camera what colour “white” looks like under your "
     "current lighting. When WB is set correctly, whites look white and skin tones look natural. "
     "When WB is wrong, photos look too orange (too warm) or too blue (too cool).")

heading2(doc, "What is Colour Temperature (Kelvin)?")
body(doc,
     "Light sources have different colour temperatures, measured in Kelvin (K). A low Kelvin "
     "number means warm, orange-tinted light (like a candle or old light bulb). A high Kelvin "
     "number means cool, blue-tinted light (like an overcast sky or shade). The camera needs to "
     "know this so it can correct the colour cast and render the scene neutrally.")

two_col_table(doc,
    ["Preset", "Colour Temp.", "Typical Lighting", "Effect on Photo"],
    [
        ["BULB",         "~2800 K", "Incandescent / tungsten bulbs",      "Removes orange cast indoors"],
        ["INDOOR",       "~3500 K", "LED or mixed indoor lighting",        "Neutral-warm indoor result"],
        ["FLUORESCENT",  "~4500 K", "Office fluorescent tubes",            "Removes greenish cast"],
        ["DAYLIGHT",     "~5500 K", "Outdoor midday — neutral reference",  "Natural, neutral colours"],
        ["CLOUDY",       "~6500 K", "Overcast sky",                        "Warms up the image slightly"],
        ["SHADE",        "~8000 K", "Open shade / blue sky reflected",     "Removes cool-blue cast"],
    ],
    col_widths=[3, 2.5, 4.5, 5.5]
)
add_spacer(doc, 12)

heading2(doc, "How to Adjust White Balance")
body(doc,
     "Tap any WB preset in the WB: row at the bottom of the camera screen. The preview colour "
     "shifts instantly so you can judge the result live. The same correction is baked into every "
     "saved photo — what you see in the preview is exactly what you get in the saved photo.")

heading2(doc, "Practical Tips")
bullet(doc,
       "Set WB to BULB (2800K) or INDOOR (3500K). The yellow cast disappears instantly in the preview.",
       bold_prefix="Indoors under a light bulb:  ")
bullet(doc,
       "Set WB to CLOUDY (6500K) to remove the cool blue cast that overcast light creates.",
       bold_prefix="Outdoors on an overcast day:  ")
bullet(doc,
       "Leave at DAYLIGHT (5500K) — this is the neutral default.",
       bold_prefix="Outdoors in bright sun:  ")
bullet(doc,
       "Deliberately set the “wrong” WB for effect — BULB (2800K) outdoors gives a "
       "dramatic cinematic blue look; SHADE (8000K) indoors adds a golden warm-hour feel.",
       bold_prefix="Creative use:  ")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 10. FLASH
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "10.  Flash")
body(doc, "Tap the FLASH icon in the controls bar to cycle through flash modes:")
bullet(doc, "Flash never fires.",                              bold_prefix="OFF  ")
bullet(doc, "Flash fires automatically when the scene is dark.", bold_prefix="AUTO  ")
bullet(doc, "Flash always fires.",                             bold_prefix="ON  ")
bullet(doc, "Continuous torch light — useful for close-up shots or video.", bold_prefix="TORCH  ")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 11. RAW CAPTURE
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "11.  RAW Capture")
body(doc,
     "Tap the RAW button in the controls bar to toggle RAW capture on or off. When RAW is on:")
bullet(doc, "The photo is saved as a DNG file — an unprocessed, lossless file.")
bullet(doc, "No colour look, bokeh, or lens simulation is applied.")
bullet(doc,
       "The file is larger but contains maximum image data for editing in Adobe Lightroom, "
       "Snapseed, or any DNG-compatible editor.")
body(doc, "RAW files are saved to Pictures/iCamera/ alongside your JPEG photos.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 12. PHOTO QUALITY
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "12.  Photo Quality Setting — STD vs HQ")
body(doc,
     "The QUAL button in the controls bar lets you choose between two capture quality levels. "
     "Tapping it toggles between STD (Standard) and HQ (High Quality).")

heading2(doc, "What is a Megapixel?")
body(doc,
     "A digital photo is made up of millions of tiny coloured dots called pixels. A megapixel (MP) "
     "is simply one million of those dots. More dots = more fine detail — but also a larger file "
     "and longer processing time.")

heading2(doc, "STD — Standard Quality (~8 MP)")
body(doc,
     "Standard mode captures at around 8 megapixels (~3264 × 2448 dots). "
     "This is the default and recommended for everyday use.")
bullet(doc, "Shutter response is near-instant — the camera feels as fast as the built-in camera app.")
bullet(doc, "Photos are saved quickly — the thumbnail appears within about one second.")
bullet(doc, "8MP is more than enough for every screen: phone, tablet, TV, Instagram, WhatsApp.")
bullet(doc, "Can print up to approximately A3 size (43 cm × 33 cm) at full quality.")
tip_box(doc, "Use STD for all everyday shooting. You will not see any quality difference on a screen.")

heading2(doc, "HQ — High Quality (Full Sensor, up to 13 MP)")
body(doc,
     "High Quality mode uses the camera sensor at its maximum resolution — up to 13 megapixels "
     "(~4208 × 3120 dots) on the Xiaomi Pad 6 and similar devices.")
bullet(doc, "More fine detail — useful for very large prints or heavy cropping.")
bullet(doc, "Processing takes longer — expect 2–4 seconds after the shutter click.")
bullet(doc, "File sizes are noticeably larger.")
tip_box(doc, "Use HQ only when you plan to print larger than A3 or need to crop significantly into the photo.")

heading2(doc, "How to Switch Quality")
body(doc,
     "In the controls bar, look for the QUAL button. The label above it shows the current setting: "
     "STD = Standard (8MP, fast)  ·  HQ = High Quality (full sensor, slower but maximum detail). "
     "Tap QUAL to toggle. The camera briefly reinitialises before the new resolution takes effect.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 13. SAVING & FINDING YOUR PHOTOS
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "13.  Saving & Finding Your Photos")
body(doc, "All photos are automatically saved to:")

para = doc.add_paragraph()
set_para_spacing(para, before=20, after=20)
para.paragraph_format.left_indent = Cm(1.5)
run = para.add_run("Pictures  /  iCamera  /")
run.font.name      = FONT_MONO
run.font.size      = Pt(10)
run.font.bold      = True
run.font.color.rgb = C_DARK

body(doc,
     "Access them via the Files app, Google Photos, or any gallery app on your device. "
     "JPEG photos are saved at 92% quality. DNG files are at full sensor resolution.")
body(doc,
     "Tap the thumbnail in the bottom-left corner of the camera screen to open your device’s "
     "gallery directly.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# 14. TIPS & TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "14.  Tips & Troubleshooting")

heading2(doc, "Tips for Better Shots")
bullet(doc,
       "Set ISO 800–1600 and shutter 1/60 to avoid blur while keeping noise under control. "
       "The CLASSIC film look’s lifted shadows hide high-ISO grain naturally.",
       bold_prefix="Low-light PRO mode:  ")
bullet(doc,
       "CLASSIC works beautifully in warm afternoon light. B&W excels in harsh midday sun where "
       "high contrast becomes a feature, not a flaw.",
       bold_prefix="Match look to light:  ")
bullet(doc,
       "The strong vignette naturally draws the eye to the centre of the frame and the bokeh "
       "separates your subject from distracting backgrounds.",
       bold_prefix="Noctilux 50mm for portraits:  ")
bullet(doc,
       "STD is fine for social media. Switch to HQ when you plan to enlarge or crop the image.",
       bold_prefix="HQ for prints:  ")
bullet(doc,
       "Find the exposure settings for your lighting, then shoot multiple frames without letting "
       "auto-exposure hunt. Consistent exposure across a series makes editing easier.",
       bold_prefix="Lock exposure in PRO mode:  ")
bullet(doc,
       "Without WB correction, overcast light photographs cold and grey. The SHADE or CLOUDY "
       "preset warms it back to a more natural, pleasing tone.",
       bold_prefix="WB for overcast days:  ")

heading2(doc, "Troubleshooting")

heading3(doc, "Preview looks too dark in PRO mode")
body(doc,
     "The viewfinder shows the real sensor output. A dark preview means the shutter speed is too "
     "fast or ISO too low for the available light. Slow down the shutter or raise ISO.")

heading3(doc, "Photos are blurry")
body(doc,
     "At shutter speeds slower than 1/60, hand movement causes blur. Rest your phone on a flat "
     "surface, or raise the shutter speed and increase ISO to compensate.")

heading3(doc, "Photos look very grainy")
body(doc,
     "ISO is set too high. Lower ISO and slow down the shutter speed instead, or move to a "
     "brighter environment.")

heading3(doc, "PRO mode scroll wheels do not appear")
body(doc,
     "Tap the PRO button in the mode bar. Two scroll wheels (SHUTTER left, ISO right) should "
     "appear between the lens selector and the mode bar. Scroll left or right and lift your "
     "finger to apply the new value.")

heading3(doc, "Camera shows an error on launch")
body(doc,
     "Ensure the app has Camera permission. Go to Settings → Apps → iCamera → "
     "Permissions → Camera → Allow.")

heading3(doc, "Photos are not saving")
body(doc,
     "Ensure the app has Storage permission. On Android 13+, this means Photos & Videos "
     "permission. Go to Settings → Apps → iCamera → Permissions → "
     "Photos & Videos → Allow.")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# FOOTER PAGE
# ════════════════════════════════════════════════════════════════════════════
add_horizontal_rule(doc, color="C80000", thickness=8)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=40, after=0)
run  = para.add_run("iCamera  ·  tcw3  ·  2026")
run.font.name      = FONT_BODY
run.font.size      = Pt(9)
run.font.color.rgb = C_SUBTLE

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(para, before=0, after=0)
run  = para.add_run("All images processed on-device. No data is uploaded to any server.")
run.font.name      = FONT_BODY
run.font.size      = Pt(8.5)
run.font.italic    = True
run.font.color.rgb = C_RULE

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
out = os.path.join(BASE, "iCamera_User_Manual_v2.docx")
doc.save(out)
print(f"Saved: {out}")
