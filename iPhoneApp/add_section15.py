#!/usr/bin/env python3
"""Append Section 15 (iPhone sensor / resolution explainer) to iCamera_User_Manual_v2.docx."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r'C:\Users\tomcw\App_tcw3\iCamera\iPhoneApp\iCamera_User_Manual_v2.docx'
doc = Document(DOC_PATH)

C_LEICA_RED  = RGBColor(0xC8, 0x00, 0x00)
C_DARK       = RGBColor(0x1A, 0x1A, 0x1A)
C_MID        = RGBColor(0x44, 0x44, 0x44)
C_SUBTLE     = RGBColor(0x66, 0x66, 0x66)
FONT_BODY    = "Georgia"
FONT_HEADING = "Georgia"

def set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'),     str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=240, after=80)
    run  = para.add_run(text.upper())
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(14)
    run.font.bold      = True
    run.font.color.rgb = C_LEICA_RED
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '4')
    btm.set(qn('w:space'), '4')
    btm.set(qn('w:color'), 'C80000')
    pBdr.append(btm)
    pPr.append(pBdr)
    return para

def heading2(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=160, after=40)
    run  = para.add_run(text)
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.color.rgb = C_DARK
    return para

def body(doc, text, indent=False):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=60, line=276)
    if indent:
        para.paragraph_format.left_indent = Cm(0.6)
    run  = para.add_run(text)
    run.font.name  = FONT_BODY
    run.font.size  = Pt(10)
    run.font.color.rgb = C_MID
    return para

def bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    set_para_spacing(para, before=0, after=40, line=260)
    para.paragraph_format.left_indent       = Cm(0.8)
    para.paragraph_format.first_line_indent = Cm(-0.4)
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = C_DARK
        r = para.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(10)
        r.font.color.rgb = C_MID
    else:
        r = para.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(10)
        r.font.color.rgb = C_MID
    return para

def tip_box(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=60, after=60)
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FFF3F3')
    pPr.append(shd)
    r = para.add_run("▶  TIP:  ")
    r.font.name      = FONT_BODY
    r.font.size      = Pt(9.5)
    r.font.bold      = True
    r.font.color.rgb = C_LEICA_RED
    r2 = para.add_run(text)
    r2.font.name     = FONT_BODY
    r2.font.size     = Pt(9.5)
    r2.font.color.rgb = C_MID
    return para

def two_col_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, '222222')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        set_para_spacing(para, before=40, after=40)
        run  = para.add_run(h)
        run.font.name      = FONT_BODY
        run.font.size      = Pt(9)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            set_para_spacing(para, before=30, after=30)
            run  = para.add_run(val)
            run.font.name      = FONT_BODY
            run.font.size      = Pt(9)
            run.font.bold      = (ci == 0)
            run.font.color.rgb = C_DARK if ci == 0 else C_MID
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return tbl

def add_spacer(doc, pts=6):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=pts)
    para.add_run("")
    return para

# ─────────────────────────────────────────────────────────────────────────────
# Add page break before new section
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 15. UNDERSTANDING IPHONE CAMERA SENSORS & RESOLUTION
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "15.  Understanding iPhone Camera Sensors & Resolution")
body(doc,
     "This section answers a question many iPhone users ask: why does a photo from iCamera on "
     "a 48 MP iPhone look the same pixel count as a photo from an iPhone 13? The answer reveals "
     "something important about how modern smartphone cameras actually work.")

heading2(doc, "Why iCamera Captures at 12 MP on Every iPhone")
body(doc,
     "iCamera uses Apple's AVCaptureSessionPresetPhoto — the same capture mode the built-in "
     "Camera app uses for standard JPEG and HEIF photos. On every current iPhone, this mode "
     "produces a 12 MP output (approximately 4032 x 3024 pixels in landscape, or 3024 x 4032 "
     "in portrait). This is true whether your iPhone has a 12 MP, 48 MP, or any future sensor.")
tip_box(doc,
        "12 MP is more than enough for any screen (phone, TV, laptop) and prints up to A2 size "
        "(59 cm x 42 cm) at full quality. Megapixel count beyond 12 MP makes almost no visible "
        "difference on any display you will use in daily life.")

heading2(doc, "What Does a 48 MP Sensor Actually Do?")
body(doc,
     "On iPhones with 48 MP sensors (iPhone 15 onwards), Apple uses a technique called "
     "pixel binning. The 48 million tiny sensor pixels are grouped into clusters of four, "
     "and each cluster is combined into one output pixel. The result is still a 12 MP photo "
     "-- but each of those 12 million pixels was built from four real sensor measurements "
     "instead of one.")

bullet(doc,
       "Four pixels worth of light data merged into one output pixel means far less noise. "
       "Your photos in dim restaurants, indoors, or at dusk are noticeably cleaner and "
       "sharper -- even though the pixel count is identical to iPhone 13.",
       bold_prefix="Better low-light quality:  ")
bullet(doc,
       "When you zoom in on a 12 MP photo from a 48 MP iPhone on screen, you will see "
       "more fine detail and less grain in dark areas compared to the same shot from a "
       "12 MP sensor iPhone.",
       bold_prefix="More detail when zooming in:  ")
bullet(doc,
       "The 48 MP sensor captures a wider tonal range -- more detail preserved in very bright "
       "highlights (sky, windows) AND very dark shadows simultaneously in the same photo.",
       bold_prefix="Wider dynamic range:  ")

heading2(doc, "Lossless 2x Zoom on 48 MP iPhones")
body(doc,
     "One of the most practical benefits of a 48 MP sensor is lossless 2x zoom. When you "
     "zoom to 2x on a 48 MP iPhone, the camera does not perform digital zoom (which degrades "
     "quality). Instead it uses only the centre crop of the 48 MP sensor -- which is exactly "
     "12 MP -- giving you a true optically-equivalent zoom with no quality loss. On an "
     "iPhone 13 (12 MP sensor), 2x zoom is always digital zoom and the photo quality "
     "noticeably degrades.")

heading2(doc, "What About Full 48 MP -- Can iCamera Access It?")
body(doc,
     "Full 48 MP capture is only available through Apple ProRAW -- a separate RAW format that "
     "bypasses the standard capture pipeline entirely. This requires a different API and produces "
     "files of 50-80 MB each. iCamera's current RAW mode saves a standard DNG via the normal "
     "capture path, which is still 12 MP. Full 48 MP ProRAW support is a potential future "
     "addition for professional users who need it.")
tip_box(doc,
        "For 95% of users -- social media, printing, sharing -- 12 MP JPEG or HEIF is the right "
        "format. ProRAW / 48 MP is for professional photographers who edit extensively in "
        "Lightroom or similar tools and need every pixel the sensor can produce.")

heading2(doc, "The Real Reasons to Upgrade Your iPhone Camera")
body(doc,
     "The megapixel number on the box is one of the least important factors in final photo "
     "quality. Here is what actually improves between iPhone generations:")

two_col_table(doc,
    ["Improvement", "What It Means in Practice"],
    [
        ["Larger physical sensor size",
         "Bigger sensor = each pixel captures more light. Night photos and indoor shots are "
         "cleaner and more detailed even at 12 MP."],
        ["Better low-light performance",
         "Newer sensors have less electronic noise at high ISO. Photos in dim light look "
         "cleaner without needing to raise ISO as high."],
        ["Wider dynamic range",
         "More detail preserved in bright highlights (sky, windows) and dark shadows "
         "in the same photo simultaneously."],
        ["Lossless 2x zoom (48 MP iPhones)",
         "True optical-quality zoom at 2x using the centre sensor crop. iPhone 13 must "
         "use digital zoom at 2x which softens the image."],
        ["Improved computational photography",
         "Newer Apple chips run better Smart HDR, Deep Fusion, and Photonic Engine "
         "algorithms that further improve sharpness and colour in every shot."],
        ["Better optics",
         "Each generation brings sharper lenses, wider apertures, and better edge-to-edge "
         "sharpness -- all of which improve the 12 MP output quality noticeably."],
    ],
    col_widths=[5.5, 10]
)
add_spacer(doc)

body(doc,
     "The bottom line: the final JPEG or HEIF photo from an iPhone 15, 16, or 17 is still "
     "12 MP -- the same pixel count as iPhone 13. But the quality of those 12 MP is "
     "meaningfully better, especially in low light, at 2x zoom, and in high-contrast scenes. "
     "You will see and feel the difference in real shooting -- just not in the file size or "
     "pixel count.")
add_spacer(doc)

doc.save(DOC_PATH)
print("Section 15 appended and saved:", DOC_PATH)
