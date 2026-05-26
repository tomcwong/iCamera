#!/usr/bin/env python3
"""Append HEIF / PRO-mode manual exposure sections to iCamera_User_Manual_v2.docx."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r'C:\Users\tomcw\App_tcw3\iCamera\iPhoneApp\iCamera_User_Manual_v2.docx'
doc = Document(DOC_PATH)

C_LEICA_RED = RGBColor(0xC8, 0x00, 0x00)
C_DARK      = RGBColor(0x1A, 0x1A, 0x1A)
C_MID       = RGBColor(0x44, 0x44, 0x44)
C_SUBTLE    = RGBColor(0x66, 0x66, 0x66)
FONT_BODY   = "Georgia"
FONT_HEADING= "Georgia"

def set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'),     str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=240, after=80)
    run  = para.add_run(text.upper())
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(14)
    run.font.bold      = True
    run.font.color.rgb = C_LEICA_RED
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '4')
    btm.set(qn('w:space'), '4')
    btm.set(qn('w:color'), 'C80000')
    pBdr.append(btm)
    pPr.append(pBdr)
    return para

def heading2(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=160, after=40)
    run  = para.add_run(text)
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.color.rgb = C_DARK
    return para

def heading3(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=100, after=30)
    run  = para.add_run(text)
    run.font.name      = FONT_HEADING
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.italic    = True
    run.font.color.rgb = C_MID
    return para

def body(doc, text, indent=False):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=60, line=276)
    if indent:
        para.paragraph_format.left_indent = Cm(0.6)
    run  = para.add_run(text)
    run.font.name      = FONT_BODY
    run.font.size      = Pt(10)
    run.font.color.rgb = C_MID
    return para

def bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    set_para_spacing(para, before=0, after=40, line=260)
    para.paragraph_format.left_indent       = Cm(0.8)
    para.paragraph_format.first_line_indent = Cm(-0.4)
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.font.name      = FONT_BODY
        r.font.size      = Pt(10)
        r.font.bold      = True
        r.font.color.rgb = C_DARK
        r = para.add_run(text)
        r.font.name      = FONT_BODY
        r.font.size      = Pt(10)
        r.font.color.rgb = C_MID
    else:
        r = para.add_run(text)
        r.font.name      = FONT_BODY
        r.font.size      = Pt(10)
        r.font.color.rgb = C_MID
    return para

def tip_box(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=60, after=60)
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FFF3F3')
    pPr.append(shd)
    r = para.add_run("▶  TIP:  ")
    r.font.name      = FONT_BODY
    r.font.size      = Pt(9.5)
    r.font.bold      = True
    r.font.color.rgb = C_LEICA_RED
    r2 = para.add_run(text)
    r2.font.name     = FONT_BODY
    r2.font.size     = Pt(9.5)
    r2.font.color.rgb = C_MID
    return para

def two_col_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, '222222')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        set_para_spacing(para, before=40, after=40)
        run  = para.add_run(h)
        run.font.name      = FONT_BODY
        run.font.size      = Pt(9)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            set_para_spacing(para, before=30, after=30)
            run  = para.add_run(val)
            run.font.name      = FONT_BODY
            run.font.size      = Pt(9)
            run.font.bold      = (ci == 0)
            run.font.color.rgb = C_DARK if ci == 0 else C_MID
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return tbl

def add_spacer(doc, pts=6):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=pts)
    para.add_run("")
    return para

# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# UPDATE: PHOTO QUALITY — STD / HQ / HEIF  (replaces Section 12 in v2)
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, "12 (Updated).  Photo Quality — STD / HQ / HEIF")
body(doc,
     "The QUAL button in the Gear panel now cycles through three capture quality levels. "
     "Each tap moves to the next option: STD → HQ → HEIF → STD. "
     "(This replaces the original two-option STD/HQ toggle described in the earlier section.)")

two_col_table(doc,
    ["Setting", "Format", "Quality", "Typical File Size", "Best For"],
    [
        ["STD",  "JPEG", "85%",        "~2–3 MB", "Everyday use — fast, small, sharp"],
        ["HQ",   "JPEG", "97%",        "~5–7 MB", "Large prints, heavy cropping"],
        ["HEIF", "HEIC", "H.265 ~90%", "~2–4 MB", "Maximum quality, smallest file (Apple ecosystem)"],
    ],
    col_widths=[2, 2, 2.5, 2.5, 6.5]
)
add_spacer(doc, 10)

heading2(doc, "STD — Standard JPEG")
body(doc, "JPEG at quality 85 — the default for all everyday shooting.")
bullet(doc, "Fast processing — thumbnail appears within about one second.")
bullet(doc, "File sizes around 2–3 MB — easy to share by message or email.")
bullet(doc, "Fine for any screen or A3 print.")
tip_box(doc, "Use STD for everyday shooting. You will not see any quality difference on a phone or TV screen.")

heading2(doc, "HQ — High Quality JPEG")
body(doc, "JPEG at quality 97 — maximum detail with minimal compression artefacts.")
bullet(doc, "More fine detail — useful for very large prints or significant cropping.")
bullet(doc, "Processing takes 2–4 seconds after the shutter click.")
bullet(doc, "File sizes noticeably larger (~5–7 MB).")
tip_box(doc, "Use HQ when you plan to print larger than A3, or need to crop deeply into the photo.")

heading2(doc, "HEIF — Apple High Efficiency Image Format")
body(doc,
     "HEIF (pronounced ‘HEEF’, file extension .heic) is Apple’s native photo format — the same "
     "format the built-in iPhone Camera app uses when you select ‘High Efficiency’ in "
     "iOS Settings → Camera → Formats. It uses H.265 (HEVC) compression, which is roughly "
     "twice as efficient as JPEG at the same visual quality.")
bullet(doc,
       "Same visual quality as HQ JPEG but the file is roughly 50% smaller. "
       "A photo that would be 6 MB as HQ JPEG is typically 2–3 MB as HEIF.",
       bold_prefix="Smaller files, same quality:  ")
bullet(doc,
       "iCamera encodes your processed pixels directly into HEIF using iOS native ImageIO "
       "(CGImageDestinationCreateWithData). There is no intermediate JPEG step. "
       "The Leica Look pipeline runs on the raw pixels first, then those pixels go straight "
       "into the HEIF encoder. One encode step = maximum quality for the format.",
       bold_prefix="Direct pixel encoding:  ")
bullet(doc,
       "HEIF files open natively in Photos.app, iCloud, Safari, and Mac Preview. "
       "Windows 11 with the HEIF codec from the Microsoft Store also supports them. "
       "Some older apps may not recognise .heic — use HQ JPEG for maximum compatibility.",
       bold_prefix="Compatibility:  ")
tip_box(doc,
        "HEIF is recommended for iPhone users who want maximum quality with the smallest possible "
        "file. Use HQ JPEG if you need to share with apps or services that do not support .heic.")

heading2(doc, "HEIF vs Leica Look — Not the Same Thing")
body(doc,
     "These are two completely separate controls and are often confused:")
bullet(doc,
       "A file compression format. Controls how image data is stored on disk — not what "
       "the image looks like. Choosing HEIF vs JPEG does not change colours, vignette, or any "
       "creative effect in your photo.",
       bold_prefix="HEIF (QUAL button):  ")
bullet(doc,
       "A creative processing pipeline (colour grade, vignette, chromatic aberration, distortion) "
       "applied to the captured pixels. Controlled by the LEICA button. Runs independently of "
       "the output format. If Leica Look is OFF, the photo is saved clean in whichever format "
       "is selected.",
       bold_prefix="Leica Look (LEICA button):  ")
add_spacer(doc)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 16: HOW iCAMERA APPLIES MANUAL EXPOSURE ON iPHONE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1(doc, "16.  How iCamera Applies Manual Exposure on iPhone")
body(doc,
     "When you select PRO mode and set a specific ISO and shutter speed, iCamera sends those exact "
     "values directly to the iPhone camera sensor using Apple’s AVFoundation framework. "
     "This section explains how that works and why it matters.")

heading2(doc, "What Happens When You Enter PRO Mode")
body(doc, "The moment you tap PRO, iCamera does three things in sequence:")
bullet(doc,
       "Locks the autofocus so the focus point stays fixed. "
       "The camera will not hunt or refocus while you compose your shot.",
       bold_prefix="1. Locks focus:  ")
bullet(doc,
       "Sends your selected ISO and shutter speed directly to the sensor "
       "using AVCaptureDevice.setExposureModeCustom — the same low-level API "
       "that Apple’s own Camera app uses for manual controls. "
       "The camera’s automatic exposure brain is switched off entirely.",
       bold_prefix="2. Sets hardware exposure:  ")
bullet(doc,
       "Waits for the sensor to confirm that the new settings have been applied "
       "(via an iOS completion callback) before allowing a capture. "
       "This ensures the photo is taken with your exact values, not a stale auto-exposure reading.",
       bold_prefix="3. Waits for confirmation:  ")

heading2(doc, "What the Live HUD Shows in AUTO vs PRO")
body(doc,
     "In AUTO mode, the SS and ISO values shown in the top bar are live readings from the sensor "
     "— the camera’s automatic exposure system updates them up to twice per second as lighting "
     "changes. What you see is what the sensor is actually using right now.")
body(doc,
     "In PRO mode, the SS and ISO values shown are the values you have chosen. "
     "The sensor has been locked to those exact values and will not change them until "
     "you scroll the wheel or switch back to AUTO.")

heading2(doc, "Why Manual Exposure Matters")
two_col_table(doc,
    ["Situation", "What to do in PRO"],
    [
        ["Candle or firelight",
         "ISO 400, SS 1/125 — freezes the flame without overexposing"],
        ["Night city street",
         "ISO 1600, SS 1/30 — rest phone on surface to avoid blur"],
        ["Bright beach in noon sun",
         "ISO 100, SS 1/2000 — clean sky, no blown-out whites"],
        ["Consistent series for editing",
         "Lock ISO/SS at correct brightness — every shot in the series is identical"],
        ["Intentional motion blur",
         "ISO 100, SS 1/15 or slower — cars blur, static elements stay sharp"],
    ],
    col_widths=[5.5, 10]
)
add_spacer(doc)

heading2(doc, "The iPhone Aperture Limitation")
body(doc,
     "Unlike a traditional camera, the iPhone lens has a fixed, non-variable aperture. "
     "The iPhone 13 wide camera is always f/1.6 — you cannot change it mechanically. "
     "iCamera’s APT mode and the f-stop value in the HUD simulate aperture "
     "for bokeh rendering purposes, but the physical lens opening is always the same. "
     "This is why the EXIF data on saved photos always shows f/1.6 regardless of the "
     "APT setting you chose in iCamera.")
tip_box(doc,
        "The APT value in iCamera controls the strength of the software bokeh blur applied "
        "at capture time, not the physical lens opening. Lower f-number = stronger blur.")
add_spacer(doc)

doc.save(DOC_PATH)
print("Sections appended and saved:", DOC_PATH)
